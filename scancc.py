""" ---------------------------------------------------------------------------
    Description
    --------------------------------------------------------------------------- 
 
    This script searches and masks PAN (aka, Primary Account Number,
    or credit card numbers, or banking card number) in
    text files, word files and excel files.
    It does not work on PDF.
    It does not work on pictures.
    It does not work if the files are encrypted or password protected.
        
    To simplify its work, this script also recursively unzip files.
    This however, has NOT been fully tested yet!

    An original copy of the modified files that are modified is put in
    a quarantine folder.
 
    Author:         JL Dupont
    Date/version:   20230102
"""


""" ---------------------------------------------------------------------------
    Imports
    ----------------------------------------------------------------------- """
import os
import fnmatch
import re
import shutil
import datetime
import openpyxl
import docx
import pdfrw
import zipfile
from itertools import filterfalse

""" ---------------------------------------------------------------------------
    Constants
    ----------------------------------------------------------------------- """
C_SCANFOLDER        = "c:\\temp\\scancc"
C_QUARANTINEFOLDER  = "c:\\quarantine\\"
C_BACKUPZIPFOLDER   = C_QUARANTINEFOLDER + "zip\\"
C_TEMPFILE          = "scancc.tmp"
C_MASK1             = 'Info'
C_MASK2             = 'Secu'


""" ---------------------------------------------------------------------------
    Global variables
    ----------------------------------------------------------------------- """
g_filecount         = 0
g_pancount          = 0
g_quarantinecount   = 0


""" ---------------------------------------------------------------------------
    Functions
    ----------------------------------------------------------------------- """

def f_scantextfile(root, file):
    """ Scan a text file for PAN and mask the PAN. Original file is
        quarantined. 
    
    Args:
        root (str): root directory
        file (str): file name
        
    Returns:
        nothing
    """
    file_path = os.path.join(root, file)
    if not os.path.exists(root) or not os.path.isfile(file_path):
        return
    try:
        with open(file_path, 'r', errors="ignore", encoding='utf-8') as f:
            text = f.read()
        cclist = f_findcc(text)
        if len(cclist) > 0:
            for cc in cclist:
                text = text.replace(cc, f_maskcc(cc))
            temp_file = os.path.join(root, C_TEMPFILE)
            with open(temp_file, 'w', encoding='utf-8') as f:
                f.write(text)
            f_quarantine(root, file)
    except Exception as e:
        print(f'An error occurred: {e}')



def f_scanexcelfile(root, file):
    """ Scan an Excel file for PAN and mask the PAN. Original file is
        quarantined. 
    
    Args:
        root (str): root directory
        file (str): file name
        
    Returns:
        nothing
    """
    file_path = os.path.join(root, file)
    if not os.path.exists(root) or not os.path.isfile(file_path):
        return
    try:
        onechange = False
        with openpyxl.load_workbook(file_path, read_only=True) as workbook:
            sheet_names = workbook.sheetnames
            for sheet_name in sheet_names:
                sheet = workbook[sheet_name]
                for row in sheet.rows:
                    for cell in row:
                        cclist = f_findcc(str(cell.value))
                        if len(cclist) > 0:            
                            for cc in cclist:
                                cell.value = str(cell.value).replace(cc, f_maskcc(cc))
                                onechange = True
        if onechange:
            temp_file = os.path.join(root, C_TEMPFILE)
            workbook.save(temp_file)
            f_quarantine(root, file)
    except Exception as e:
        print(f'An error occurred: {e}')



def f_scanpdffile(root, file):
    """ Scan a PDF file for PAN and mask the PAN. Original file is
        quarantined. 
    
    Args:
        root (str): root directory
        file (str): file name
        
    Returns:
        nothing
    """
    file_path = os.path.join(root, file)
    if not os.path.exists(root) or not os.path.isfile(file_path):
        return
    try:
        onechange = False
        # Open the PDF in read-binary mode
        mypdf = pdfrw.PdfReader(file_path)
        for page in mypdf.pages:
            # Get the page contents
            contents = page.get('/Contents')
            # Decode the contents (if necessary)
            if isinstance(contents, pdfrw.objects.pdfstring.PdfString):
                contents = contents.decode()
            # Replace the text
            cclist = f_findcc(contents)
            if len(cclist) > 0:            
                for cc in cclist:
                    contents = contents.replace(cc, f_maskcc(cc))
                    onechange = True
            # Encode the contents (if necessary)
            if isinstance(page['/Contents'], pdfrw.objects.pdfstring.PdfString):
                contents = pdfrw.objects.pdfstring.PdfString(contents)
            # Update the page contents
            page['/Contents'] = contents
        if onechange:
            temp_file = os.path.join(root, C_TEMPFILE)
            pdfrw.PdfWriter().write(temp_file, mypdf)
            f_quarantine(root, file)
    except Exception as e:
        print(f'An error occurred: {e}')



def f_scanwordfile(root, file):
    """ Scan a Word file for PAN and mask the PAN. Original file is
        quarantined. 
    
    Args:
        root (str): root directory
        file (str): file name
        
    Returns:
        nothing
    """
    file_path = os.path.join(root, file)
    if not os.path.exists(root) or not os.path.isfile(file_path):
        return
    try:
        onechange = False
        document = docx.Document(file_path)
        for paragraph in document.paragraphs:
            cclist = f_findcc(str(paragraph.text))
            if len(cclist) > 0:            
                for cc in cclist:
                    paragraph.text = paragraph.text.replace(cc, f_maskcc(cc))
                    onechange = True
        for table in document.tables:
            for rows in table.rows:       
                for cell in rows.cells:
                    cclist = f_findcc(str(cell.text))
                    if len(cclist) > 0:            
                        for cc in cclist:
                            cell.text = str(cell.text).replace(cc, f_maskcc(cc))
                            onechange = True
        if onechange:
            temp_file = os.path.join(root, C_TEMPFILE)
            document.save(temp_file)
            f_quarantine(root, file)
    except Exception as e:
        print(f'An error occurred: {e}')



def f_quarantine(root, file):
    """ Quarantine a file by moving it to the quarantine folder. The file in
        the original folder is then replaced by the temporary file
        (where the PAN is masked)
    
    Args:
        root (str): root directory
        file (str): file name
        
    Returns:
        nothing
    """
    global g_quarantinecount

    g_quarantinecount += 1
    dest = os.path.join(C_QUARANTINEFOLDER, g_timestamp, root[2:], file)
    file_path = os.path.join(root, file)
    temp_file = os.path.join(root, C_TEMPFILE)
    if not os.path.exists(root) or not os.path.isfile(file_path) or not os.path.isfile(temp_file):
        return
    try:
        filetime = os.stat(file_path).st_mtime
        os.utime(temp_file, (filetime, filetime))
        # Create the destination directory if it does not exist
        if not os.path.isdir(os.path.dirname(dest)):
            os.makedirs(os.path.dirname(dest))
        # Move the file to the destination directory
        shutil.move(file_path, dest)
        # Replace with the temporary file
        shutil.move(temp_file, file_path)
    except Exception as e:
        print(f'An error occurred: {e}')



def f_zipbackup(root, file):
    """ Backup a zip file before deletion
    
    Args:
        root (str): root directory
        file (str): file name
        
    Returns:
        nothing
    """
    file_path = os.path.join(root, file)
    if not os.path.exists(root) or not os.path.isfile(file_path):
        return
    dest = os.path.join(C_, g_timestamp, root[2:], file)
    try:
        # Create the destination directory if it does not exist
        if not os.path.exists(os.path.dirname(dest)):
            os.makedirs(os.path.dirname(dest))
        # Move the file to the destination directory
        shutil.move(file_path, dest)
    except Exception as e:
        print(f'An error occurred: {e}')




def f_maskcc(card_number):
    """ Mask a PAN
    
    Args:
        card_number (str): a string containing a card number 
        
    Returns:
        a string where the card number is partially masked.
        The first 4 and last 4 digits are kept.
        The original string format is kept.
        
    """
    
    global g_pancount
    
    g_pancount += 1
    spacer = ""
    if " " in card_number:
        spacer = " "
    if "-" in card_number:
        spacer = "-"
    return  card_number[0:4] + spacer + C_MASK1 + spacer + C_MASK2 + spacer + card_number[-4:]
    

def f_luhn(card_number):
    """ Validate a card_number as a valid PAN
        by checking the first digit and applying the luhn
        algorithm
    
    Args:
        card_number (str): a string containing a card number 
        
    Returns:
        boolean.
    """

    if not card_number.startswith(("3", "4", "5", "6")):
        return False


    # Remove # and ' '
    card_number = "".join(filterfalse(lambda x: x in " -", card_number))

    # Reverse the card number
    card_number = card_number[::-1]

    # Convert to integer and double every other number
    doubled_digits = []
    for i, digit in enumerate(card_number):
        digit = int(digit)
        if i % 2 == 0:
            doubled_digits.append(digit)
        else:
            doubled_digits.append(digit * 2)

    # Subtract 9 from numbers over 9
    for i, digit in enumerate(doubled_digits):
        if digit > 9:
            doubled_digits[i] = digit - 9

    # Sum all digits
    sum_of_digits = sum(doubled_digits)

    # Return validation result
    return sum_of_digits % 10 == 0


def f_findcc(text):
    """ Search for all PAN in a string
    
    Args:
        text (str): text to search in
        
    Returns:
        a list of string: each string is a validated PAN
    """
    cc_pattern = r'(?:\d[ -]*?){16}'
    cc = re.findall(cc_pattern, text)
    
    i = 0
    while i < len(cc):
        if f_luhn(cc[i]) == False:
            cc.pop(i)
        else:
            i += 1    
    return cc
    
def f_testfile(root, file):
    """ Check is a file is a valid candidate for 
        PAN search and mask. If it is, scan the file.
    
    Args:
        root (str): root directory
        file (str): file name
        
    Returns:
        nothing
    """
    if f_isfiletext(os.path.join(root, file)):
        f_scantextfile(root, file)
    if f_isfileexcel(os.path.join(root, file)):
        f_scanexcelfile(root, file)
    if f_isfileword(os.path.join(root, file)):
        f_scanwordfile(root, file)
    """ does not work 
    if f_isfilepdf(os.path.join(root, file)):
        f_scanpdffile(root, file)
    """

def f_isfileword(file_path):
    """ Check is a file is an MS Word document.
        Checking both the extension and the format.
    
    Args:
        root (str): root directory
        file (str): file name
        
    Returns:
        Boolean
    """

    # Check the file extension
    if os.path.splitext(file_path)[1] in ['.doc','.docx']:
        is_word_file = True
        try:
            document = docx.Document(file_path)
        except:
            is_word_file = False
    else:
        is_word_file = False
    return is_word_file

def f_isfileexcel(file_path):
    """ Check is a file is an MS Excel document.
        Checking both the extension and the format.
    
    Args:
        root (str): root directory
        file (str): file name
        
    Returns:
        Boolean
    """

    # Check the file extension
    if os.path.splitext(file_path)[1] in ['.xlsx', '.xlsm', '.xltx', '.xltm']:
        is_excel_file = True
        try:
            workbook = openpyxl.load_workbook(file_path)
        except:
            # If the file cannot be opened as an Excel workbook, it is not an Excel file
            is_excel_file = False
    else:
        # If the file does not have a recognized Excel file extension, it is not an Excel file
        is_excel_file = False
    return is_excel_file


def f_isfiletext(file_path):
    """ Check is a file is a text file, 
        regardless of the extension.
    
    Args:
        root (str): root directory
        file (str): file name
        
    Returns:
        Boolean
    """

    # Check if the file exists
    if not os.path.exists(file_path):
        return False

    file_extension = os.path.splitext(file_path)[1]
    if file_extension.lower() == ".txt":
        return True

    # Check if the file is a text file
    try:
        with open(file_path, 'r') as f:
            try:
                f.read()
            except:
                return False
    except:
        return False
    return True
    
 
def f_isfilepdf(file_path):
    """ Check is a file is an PDF document.
        Checking both the extension and the format.
    
    Args:
        root (str): root directory
        file (str): file name
        
    Returns:
        Boolean
    """

    if not os.path.exists(file_path):
        return False


    file_extension = os.path.splitext(file_path)[1]
    if file_extension.lower() == ".pdf":
        ispdf = True
 
    # Open the file in binary mode
    with open(file_path, 'rb') as file:
        # Read the first 4 bytes of the file
        file_header = file.read(4)
        # Check if the file's header matches the PDF header
        ispdf = (file_header == b'%PDF')
    return ispdf

def f_lookforfiles(directory):
    """ Recursive function to traverse the directory tree and 
        get a list of all the files and directories
   
    Args:
        directory (str): directory to start the scan from
        
    Returns:
        Nothing
    """
    global g_filecount

    file_list = []
    for root, dirs, files in os.walk(directory):
        for file in files:
            file_list.append(os.path.join(root, file))
            g_filecount += 1
            f_testfile(root, file)  
        for dir in dirs:
            file_list.append(os.path.join(root, dir))
    return file_list


def f_lookforzip(directory):
    """ Recursive function to traverse the directory tree and 
        get a list of all the files and directories
   
    Args:
        directory (str): directory to start the scan from
        
    Returns:
        Nothing
    """


    file_list = []
    for root, dirs, files in os.walk(directory):
        for file in files:
            if file.endswith('.zip') or file.endswith('.7z'):
                f_unziprecursive(os.path.join(root, file))
                file_list.append(os.path.join(root, file))
        for dir in dirs:
            file_list.append(os.path.join(root, dir))
    return file_list


def f_unzip(zip_file, destination_folder):
    """ Extract a zip file. Make a backup beforehand.
        Delete the original zip file
   
    Args:
        zip_file (str): the file extract
        destination_folder: the destination directory
        
    Returns:
        Nothing
    
    with zipfile.ZipFile(zip_file, 'r') as zip_ref:
        zip_ref.extractall(destination_folder)
    os.remove(zip_file)
    """
    print(zip_file + " " + destination_folder)

def f_unziprecursive(zip_file, destination_folder):
    extract_zip(zip_file, destination_folder)
    for root, dirs, files in os.walk(destination_folder):
        for file in files:
            if file.endswith('.zip') or file.endswith('.7z'):
                f_unzip(os.path.join(root, file), root)
        for dir in dirs:
            f_unziprecursive(os.path.join(root, dir), os.path.join(destination_folder, dir))


""" ---------------------------------------------------------------------------
    Main
    ----------------------------------------------------------------------- """

# Get a time stamp that will be used in the quarantine and in the zip backup folder
now = datetime.datetime.now()
g_timestamp = now.strftime("%Y-%m-%d %HH%M")

# Search for zip files and extract them
f_lookforzip(C_SCANFOLDER)

# Search for files and masks PAN where applicable
f_lookforfiles(C_SCANFOLDER)


print("Number of files              : " + str(g_filecount))
print("Number of redacted files     : " + str(g_quarantinecount))
print("Number of credit card numbers: " + str(g_pancount))



""" ---------------------------------------------------------------------------
    End of file
    ----------------------------------------------------------------------- """

